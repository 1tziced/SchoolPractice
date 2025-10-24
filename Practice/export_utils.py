from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import simpleSplit
from datetime import datetime
import io
import os


def create_student_certificate(student, group_name):
    """Создание справки об обучении студента из Word-шаблона"""

    # Путь к шаблону
    template_path = 'certificate_template.docx'

    # ЗАГРУЖАЕМ ШАБЛОН (НЕ СОЗДАЕМ НОВЫЙ ДОКУМЕНТ!)
    doc = Document(template_path)

    # Подготовка данных для замены плейсхолдеров
    replacements = {
        '{{ref_number}}': f'{student.id}-СТ/{datetime.now().year}',
        '{{student_fio}}': f'{student.surname} {student.name}',
        '{{group_name}}': group_name or 'Не указана',
        '{{study_year}}': str(datetime.now().year - 2),
        '{{email}}': student.email or 'Не указан',
        '{{phone}}': student.phone or 'Не указан',
        '{{issue_date}}': datetime.now().strftime('%d.%m.%Y')
    }

    # Замена плейсхолдеров в параграфах
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Замена плейсхолдеров в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, value)

    # === ДОБАВЛЕНИЕ ПЕЧАТИ ===
    # Ищем параграф с текстом "М.П." или добавляем печать в конец документа
    stamp_path = 'stamp.png'  # Путь к файлу печати

    if os.path.exists(stamp_path):
        # Ищем место для вставки печати (обычно это последние параграфы с подписями)
        # Добавляем печать перед последним параграфом
        last_paragraph = doc.paragraphs[-1]

        # Создаем новый параграф для печати
        stamp_paragraph = last_paragraph.insert_paragraph_before()
        stamp_run = stamp_paragraph.add_run()

        # Вставляем изображение печати
        # Размер: 4см x 4см (компактная печать)
        stamp_run.add_picture(stamp_path, width=Cm(4), height=Cm(4))

        # Выравнивание по левому краю (где обычно ставится М.П.)
        stamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Добавляем отступ слева для выравнивания с текстом "М.П."
        stamp_paragraph.paragraph_format.left_indent = Cm(0.5)

    # Сохранение в буфер
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def create_schedule_excel(group, schedules, subjects_dict):
    """Создание расписания в Excel с профессиональным оформлением"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Расписание"

    # === СТИЛИ ===
    header_font = Font(name='Arial', size=14, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
    cell_font = Font(name='Arial', size=11)
    title_font = Font(name='Arial', size=16, bold=True)
    subtitle_font = Font(name='Arial', size=12, bold=True)

    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # === ШАПКА ДОКУМЕНТА ===
    ws.merge_cells('A1:G1')
    cell = ws['A1']
    cell.value = 'ГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ'
    cell.font = title_font
    cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 25

    ws.merge_cells('A2:G2')
    cell = ws['A2']
    cell.value = '"ТЕХНИЧЕСКИЙ КОЛЛЕДЖ"'
    cell.font = Font(name='Arial', size=14, bold=True)
    cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[2].height = 20

    ws.merge_cells('A3:G3')
    cell = ws['A3']
    cell.value = 'г. Москва, ул. Профессиональная, д. 15'
    cell.font = Font(name='Arial', size=10)
    cell.alignment = Alignment(horizontal='center', vertical='center')

    ws.merge_cells('A5:G5')
    cell = ws['A5']
    cell.value = f'РАСПИСАНИЕ ЗАНЯТИЙ'
    cell.font = Font(name='Arial', size=14, bold=True)
    cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[5].height = 25

    ws.merge_cells('A6:G6')
    cell = ws['A6']
    cell.value = f'Группа: {group.name}'
    cell.font = subtitle_font
    cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[6].height = 20

    ws.merge_cells('A7:G7')
    cell = ws['A7']
    cell.value = f'Учебный год: {datetime.now().year}-{datetime.now().year + 1}'
    cell.font = Font(name='Arial', size=10, italic=True)
    cell.alignment = Alignment(horizontal='center', vertical='center')

    # === ЗАГОЛОВКИ ТАБЛИЦЫ ===
    days = ['Пара', 'Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота']
    row = 9

    for col, day in enumerate(days, start=1):
        cell = ws.cell(row=row, column=col)
        cell.value = day
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = border

    ws.row_dimensions[row].height = 30

    # === НАСТРОЙКА ШИРИНЫ КОЛОНОК ===
    ws.column_dimensions['A'].width = 10
    for col in ['B', 'C', 'D', 'E', 'F', 'G']:
        ws.column_dimensions[col].width = 22

    # === ЗАПОЛНЕНИЕ РАСПИСАНИЯ ===
    lesson_times = [
        '08:30 - 10:00',
        '10:10 - 11:40',
        '12:00 - 13:30',
        '13:40 - 15:10'
    ]

    for lesson_num in range(1, 5):
        row += 1

        # Номер пары с временем
        cell = ws.cell(row=row, column=1)
        cell.value = f'{lesson_num}\n{lesson_times[lesson_num - 1]}'
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = border
        cell.font = Font(name='Arial', size=10, bold=True)

        for col, day in enumerate(days[1:], start=2):
            cell = ws.cell(row=row, column=col)

            # Поиск занятия
            schedule_item = next((s for s in schedules
                                  if s.lesson_number == lesson_num and s.day_of_week == day), None)

            if schedule_item:
                subject_name = subjects_dict.get(schedule_item.subject_id, 'Неизвестно')
                room = f'\n(ауд. {schedule_item.room})' if schedule_item.room else ''
                cell.value = f'{subject_name}{room}'
                cell.fill = PatternFill(start_color='E7E6E6', end_color='E7E6E6', fill_type='solid')
            else:
                cell.value = '-'

            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = border
            cell.font = cell_font

        ws.row_dimensions[row].height = 50

    # === ПОДПИСЬ И ДАТА ===
    row += 3
    ws.merge_cells(f'A{row}:G{row}')
    cell = ws[f'A{row}']
    cell.value = f'Дата формирования расписания: {datetime.now().strftime("%d.%m.%Y")}'
    cell.font = Font(name='Arial', size=10, italic=True)
    cell.alignment = Alignment(horizontal='center')

    row += 2
    ws.merge_cells(f'A{row}:D{row}')
    cell = ws[f'A{row}']
    cell.value = 'Заместитель директора по УР:'
    cell.font = Font(name='Arial', size=11)

    ws.merge_cells(f'E{row}:F{row}')
    cell = ws[f'E{row}']
    cell.value = '_________________'
    cell.font = Font(name='Arial', size=11)
    cell.alignment = Alignment(horizontal='center')

    cell = ws[f'G{row}']
    cell.value = 'С.П. Сидоров'
    cell.font = Font(name='Arial', size=11, bold=True)

    # Сохранение в буфер
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    return buffer


def create_student_certificate_pdf(student, group_name):
    """Создание справки студента в PDF - отличается от Word версии"""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Регистрация русского шрифта (DejaVu Sans поддерживает кириллицу)
    try:
        # Пробуем загрузить DejaVu Sans
        pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
        pdfmetrics.registerFont(TTFont('DejaVu-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
        font_regular = 'DejaVu'
        font_bold = 'DejaVu-Bold'
    except:
        # Если не получилось, пробуем Arial (Windows)
        try:
            pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
            pdfmetrics.registerFont(TTFont('Arial-Bold', 'arialbd.ttf'))
            font_regular = 'Arial'
            font_bold = 'Arial-Bold'
        except:
            # В крайнем случае используем встроенный Helvetica (но он не поддерживает кириллицу)
            font_regular = 'Helvetica'
            font_bold = 'Helvetica-Bold'

    # === РАМКА ДОКУМЕНТА ===
    c.setStrokeColorRGB(0.2, 0.2, 0.8)
    c.setLineWidth(2)
    c.rect(40, 40, width - 80, height - 80, stroke=1, fill=0)

    # === ШАПКА ===
    c.setFont(font_bold, 16)
    c.drawCentredString(width / 2, height - 80, "МИНИСТЕРСТВО ОБРАЗОВАНИЯ")

    c.setFont(font_bold, 14)
    c.drawCentredString(width / 2, height - 105, "ГОСУДАРСТВЕННОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ")
    c.drawCentredString(width / 2, height - 125, '"ТЕХНИЧЕСКИЙ КОЛЛЕДЖ"')

    # Линия разделителя
    c.setStrokeColorRGB(0.5, 0.5, 0.5)
    c.setLineWidth(1)
    c.line(80, height - 140, width - 80, height - 140)

    # === ЗАГОЛОВОК ===
    c.setFont(font_bold, 20)
    c.drawCentredString(width / 2, height - 180, "СПРАВКА")

    c.setFont(font_bold, 12)
    ref_number = f"№ {student.id}-PDF/{datetime.now().year}"
    c.drawCentredString(width / 2, height - 205, ref_number)

    # === ОСНОВНОЙ ТЕКСТ ===
    c.setFont(font_regular, 12)
    y = height - 250

    # Многострочный текст с переносами
    text_width = width - 160

    text1 = f"Выдана студенту(ке) {student.surname} {student.name}"
    c.drawString(80, y, text1)
    y -= 25

    text2 = "в том, что он(а) обучается в"
    c.drawString(80, y, text2)
    y -= 20

    c.setFont(font_bold, 12)
    text3 = 'Государственном образовательном учреждении'
    c.drawString(80, y, text3)
    y -= 20

    text4 = '"Технический колледж"'
    c.drawString(80, y, text4)
    y -= 25

    c.setFont(font_regular, 12)
    text5 = f"по программе среднего профессионального образования"
    c.drawString(80, y, text5)
    y -= 20

    text6 = f"в группе: {group_name or 'Не указана'}"
    c.drawString(80, y, text6)
    y -= 20

    text7 = f"с {datetime.now().year - 2} года по настоящее время."
    c.drawString(80, y, text7)

    # === ТАБЛИЦА С ДАННЫМИ ===
    y -= 40
    c.setFont(font_bold, 11)
    c.drawString(80, y, "Контактные данные студента:")
    y -= 25

    # Рисуем таблицу
    c.setStrokeColorRGB(0.3, 0.3, 0.3)
    c.setLineWidth(0.5)

    # Заголовки
    c.setFont(font_bold, 10)
    c.drawString(90, y, "ФИО:")
    c.drawString(90, y - 20, "Группа:")
    c.drawString(90, y - 40, "Email:")
    c.drawString(90, y - 60, "Телефон:")

    # Значения
    c.setFont(font_regular, 10)
    c.drawString(200, y, f"{student.surname} {student.name}")
    c.drawString(200, y - 20, f"{group_name or 'Не указана'}")
    c.drawString(200, y - 40, f"{student.email or 'Не указан'}")
    c.drawString(200, y - 60, f"{student.phone or 'Не указан'}")

    # Рамка таблицы
    c.rect(85, y - 75, width - 170, 90, stroke=1, fill=0)

    # === НАЗНАЧЕНИЕ ===
    y -= 100
    c.setFont(font_regular, 11)
    c.drawString(80, y, "Справка дана для предъявления по месту требования.")

    # === ДАТА ===
    y -= 40
    c.setFont(font_bold, 11)
    date_str = f"Дата выдачи: {datetime.now().strftime('%d.%m.%Y')}"
    c.drawString(80, y, date_str)

    # === ПОДПИСИ ===
    y -= 60
    c.setFont(font_regular, 11)
    c.drawString(80, y, "Директор колледжа")
    c.drawString(270, y, "_________________")
    c.setFont(font_bold, 11)
    c.drawString(420, y, "И.И. Иванов")

    y -= 30
    c.setFont(font_regular, 11)
    c.drawString(80, y, "Зам. директора по УР")
    c.drawString(270, y, "_________________")
    c.setFont(font_bold, 11)
    c.drawString(420, y, "С.П. Сидоров")

    # === ПЕЧАТЬ ===
    y -= 40
    c.setFont(font_regular, 10)
    c.drawString(80, y, "М.П.")

    # Добавляем изображение печати
    stamp_path = 'stamp.png'
    if os.path.exists(stamp_path):
        # Вставляем печать рядом с "М.П."
        c.drawImage(stamp_path, 110, y - 50, width=80, height=80, preserveAspectRatio=True, mask='auto')

    c.save()
    buffer.seek(0)

    return buffer